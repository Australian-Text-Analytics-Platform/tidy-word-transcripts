# ---
# jupyter:
#   jupytext:
#     cell_metadata_filter: -all
#     formats: py:percent,ipynb
#     text_representation:
#       extension: .py
#       format_name: percent
#       format_version: '1.3'
#       jupytext_version: 1.19.1
#   kernelspec:
#     display_name: Python 3 (ipykernel)
#     language: python
#     name: python3
# ---

# %% [markdown]
# # The Transcript Compiler
#
# This notebook will assist you in creating a computationally consistent view of
# transcripts of spoken materials that are recorded in Word documents (.docx). It will
# help you do things like:
#
# 1. Extract speaker codes.
# 2. Extract transcribed text.
# 3. Add further information about who's speaking
# 4. Add further information about the context of your transcripts (eg. if they're from
#    different parts of a multi phase study).
# 4. Identify and correct inconsistencies in formatting and speaker information.
# 6. Identify possible quality issues, such as missing speaker codes.
# 5. Connect your transcripts to their associated audio recordings.
# 6. Segment your transcripts into different components with labels - for example to
#    break up a semi-structured interview into topical segments for comparison.
#
# Apart from providing some structure for entering consistent information in and about
# your transcripts, this is primarily intended to help you get the most out of the
# transcripts you already have by allowing you to use them with computational tools for
# searching and filtering.
#
# Note that this tool assumes you already have transcripts in Word format (.docx) - if
# you are just starting out transcribing audio you may want to consider other tools and
# formats for transcribing that solve many of the issues we attempt to address here. We
# also assume that these Word transcripts you have will continue to be the version of
# record for your transcripts: the spreadsheet we compile here does not replace these
# but instead complements them.
#
# TODO: Working with Word files or PDF's that aren't transcripts - checkout our document
# text extractor.
# TODO: link to guidance on transcribing speech.


# %% [markdown]
# # Word Transcript Conventions and How We Compile Them
#
# This notebook relies on a common set of conventions that we have observed in many Word
# based transcripts: because Word is a document preparation tool and not a data tool
# you may need to make adjustments to your Word files or change the configuration of
# this notebook. We aim to have clear and transparent failure modes, so even if things
# aren't quite right it can still be useful.
#
# We principally rely on two common conventions:
#
# 1. A new paragraph in Word is a new turn in the transcript.
# 2. The speaker code is separated from the text of what they said by colon-tab `:  `.
#
# Download the <a href="../example_transcript/transcript_format_example.docx"
# download="">annotated example transcript</a> to understand the expected format.
#
# What doesn't work:
#
# - If line-breaks are manually inserted to wrap text.
# - If you include non transcript material in your transcripts such as headers.
# - If your speaker codes aren't consistently marked with the same punctuation.
# - Information only present in styles like *bold* or _italics_ are ignored.
# - If your transcript lines are in tables.

# %% [markdown]
# # Recommended Workflow
#
# TODO: Make a diagram for this not just another list.
#
# 1. Start with whatever Word documents you have and upload them.
# 2. Download the output spreadsheet and examine the different sheets.
# 3. If major consistency problems are evident, identify and fix worst parts.
# 4. When major problems with transcripts are fixed, move onto entering metadata.


# %% [markdown]
#
# # Upload your Transcripts
#
# Put your transcripts in a zip file and upload using the button below.
#
# Or: upload a set of docx files and an associated spreadsheet.
#
# If you upload a new zip file, it will completely replace all existing files.
#
# If you're running this for the first time with these transcripts, don't worry about
# this.
#
# Optionally, include the spreadsheet created from an earlier run of the tool in the zip
# file. The contents of that spreadsheet will be merged into the output of rerunning
# this tool so you don't have to enter any information again.
#

# %%
import ipywidgets

uploaded_via_jupyter = False

uploader = ipywidgets.FileUpload(accept=".docx", multiple=True)
display(uploader)


# %% [markdown]
#
# # Step 2 - Run This Cell
#
#

# %%
import glob
from pathlib import Path

from docx import Document

# Find all the uploaded .docx files and process them one by one

from docx import Document

transcripts_folder = Path("uploaded_transcripts")

transcript_rows = []

# This looks for colon-tab (the \t is a tab character) as the speaker-text separator.
# If you use a different convention you can try changing this.
speaker_text_separator = ":\t"

for transcript_filepath in glob.glob(str(transcripts_folder / "**.docx")):

    filename = Path(transcript_filepath).relative_to(transcripts_folder)

    print("Processing", filename)

    with open(transcript_filepath, mode="rb") as transcript_file:

        # Load the word document
        doc = Document(transcript_file)

        para_no = 1
        segment_no = 1

        # We're going to treat each paragraph in the file as a single speaker's turn.
        # This is why headers need to be wrapped in a table or similar, otherwise we
        # end up the header included as part of the transcript.
        for paragraph in doc.paragraphs:

            # The text of this paragraph
            para_text = paragraph.text

            # Handle blank lines.
            # If they're at the start of the file just ignore them, otherwise use them
            # as segment boundaries within the transcript.
            if not para_text.strip():
                if para_no == 1:
                    continue

                segment_no += 1
                continue

            # Identify speakers by looking for the first colon character then a tab.
            # If colon-tab is not matched, no speaker will be assigned to this turn.
            speaker_code, sep, text = para_text.partition(speaker_text_separator)

            # If nothing in the line matches the speaker/text separator, keep the whole
            # line, and don't record a speaker
            if sep == "":
                text = para_text
                speaker_code = None

            transcript_rows.append(
                (str(filename), para_no, speaker_code, text, segment_no)
            )

            para_no += 1

print(transcript_rows)

# %% [markdown]
# # Finding and fixing issues
#
# 1. Identify and fix speaker code issues first.
#   - One off mispellings: fix in the source file and re-upload that file.
#   - Systematic inconsistencies: one transcript may have Interviewer, the other
#   - may have interviewer (lowercase).
# 2. Then add speaker information as relevant. This will vary depending on your

# %%
from collections import Counter
from itertools import groupby


def extract_transcript_info(transcript_rows):
    """
    Extract a summary of the extracted transcript rows for each transcript.

    """

    by_transcript = groupby(transcript_rows, key=lambda x: x[0])

    for transcript, rows in by_transcript:
        all_rows = list(rows)
        print(transcript, all_rows)


def extract_segment_info(transcript_rows):
    """
    Extract a summary of the extracted segments for each transcript.

    """
    by_segment = groupby(transcript_rows, key=lambda x: (x[0], x[-1]))

    for segment, rows in by_segment:
        all_rows = list(rows)
        print(segment, all_rows)


def extract_speaker_code_info(transcript_rows):
    """
    Extract a summary of the extracted speaker codes for each transcript.

    """
    key = lambda x: (x[0], x[2])
    by_speaker = groupby(sorted(transcript_rows, key=key), key=key)

    for speaker, rows in by_speaker:
        all_rows = list(rows)
        print(speaker, all_rows)


extract_transcript_info(transcript_rows)
extract_segment_info(transcript_rows)
extract_speaker_code_info(transcript_rows)


for filename, para_no, speaker_code, text, segment_no in transcript_rows:
    pass

# %% Write the output file


# %% [markdown]
# # Pre-process conversations
#
# This step will aim to extract each turn of each uploaded conversation, and separate
# the speaker from the turn content. This step will only be as consistent as your
# transcripts are.
#
# This will completely delete and recreate the state of processed conversations - if
# you want to upload new or edited files, make the contents of the conversations folder
# match what you want, and re-run this cell.

# %%
# Start by setting up a small database to hold the processed information.

import sqlite3

convo_db_path = conversations_path / "conversations.db"

convo_db = sqlite3.connect(convo_db_path, isolation_level=None)

convo_db.executescript("""
    DROP table if exists turn;
    CREATE table turn (
        turn_id integer primary key,
        source_file,
        turn_no,
        speaker,
        turn,
        unique(source_file, turn_no)
    )

    """)

# Then we'll identify and load conversation turns from each transcript.
import glob

from docx import Document

convo_db.execute("begin")

for transcript_filepath in glob.glob(str(conversations_path / "**.docx")):
    filename = Path(transcript_filepath).relative_to(conversations_path)
    print("Processing", transcript_filepath)

    with open(transcript_filepath, "rb") as transcript_file:

        # Load the word document
        doc = Document(transcript_file)

        # We're going to treat each paragraph in the word file as a turn
        for turn_no, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text

            # Identify speakers by looking for the colon character then a tab.
            # If colon-tab is not matched, no speaker will be assigned to this turn.
            speaker_split = para_text.split(":\t")

            if len(speaker_split) == 2:
                speaker, text = speaker_split
            else:
                speaker, text = None, para_text

            convo_db.execute(
                "INSERT into turn values (?, ?, ?, ?, ?)",
                (None, transcript_filepath, turn_no, speaker, text),
            )

convo_db.execute("commit")
