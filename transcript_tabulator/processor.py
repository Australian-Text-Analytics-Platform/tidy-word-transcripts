"""
Processes transcripts in Word docs to a tidy form of transcripts in xlsx.

"""

import dataclasses as dc
import re
import typing
from collections import Counter
from io import BytesIO

from openpyxl import Workbook, load_workbook

from docx import Document


def extract_turns(transcript_doc: Document, split_speaker_on: re.Pattern):
    """
    Extract the turns from the provided docx Document.

    split_speaker_on is a regular expression Pattern to split the turn into speaker and
    remaining components. Only the first match is split.

    If no match is found to split on, the speaker will be marked as absent and the
    full text of the turn used as the spoken text.

    """

    segment_no = 1
    last_para_has_content = False

    # We're going to treat each paragraph in the word file as a turn
    for turn_no, paragraph in enumerate(transcript_doc.paragraphs):

        # Note this is where we discard styling information
        para_text = paragraph.text

        # Check for a segment break in the transcript.
        # Make sure to merge multiple empty paragraphs into one segment break.
        if not para_text and last_para_has_content:
            segment_no += 1
            last_para_has_content = False
            continue

        if para_text:
            last_para_has_content = True

        # Use the provided splitter and the first place it matches to separate the
        # text into speaker and turn.
        potential_split = split_speaker_on.split(para_text, 1)

        if len(potential_split) == 2:
            speaker_code, text = potential_split
        else:
            speaker_code, text = "", para_text

        # turn_no + 1 so we count from 1 like most people expect.
        yield [segment_no, turn_no + 1, speaker_code, text]


@dc.dataclass
class Turn:
    source_file: str
    segment_no: int
    turn_no: int
    speaker_code: str
    transcription: str


@dc.dataclass
class RowWithExtraFields:
    """
    Container for rows that might contain extra information that needs to merged in.

    The inheriting class needs to have an extra_fields attribute as a mapping holding
    keys and values of the extra fields.

    extra_fields is assumed to be consistent by having the same name, number, and
    ordering of fields within the processing of a single table of rows in this
    process.

    """

    def __post_init__(self):

        self.required_fields = [
            f.name for f in dc.fields(self) if f.name != "extra_fields"
        ]
        self.extra_field_names = list(self.extra_fields)

    def as_header_row(self):

        extras = list(self.extra_fields)

        return [*self.required_fields, *self.extra_field_names]

    def as_row(self):

        named = dc.asdict(self)

        return [
            *(named[f] for f in self.required_fields),
            *(self.extra_fields[f] for f in self.extra_field_names),
        ]

    def merge_extras(self, from_instance):

        self.extra_fields = from_instance.extra_fields
        self.__post_init__()


@dc.dataclass
class Segment(RowWithExtraFields):
    source_file: str
    segment_no: int
    turn_count: typing.Optional[int] = 0
    segment_name: typing.Optional[str] = ""
    extra_fields: typing.Optional[dict[str, typing.Any]] = dc.field(
        default_factory=dict
    )

    def as_key(self):
        return (self.source_file, self.segment_no)


@dc.dataclass
class SpeakerCode(RowWithExtraFields):
    source_file: str
    speaker_code: str
    turn_count: typing.Optional[int] = 0
    extra_fields: typing.Optional[dict[str, typing.Any]] = dc.field(
        default_factory=dict
    )

    def as_key(self):
        return (self.source_file, self.speaker_code)


@dc.dataclass
class Transcript(RowWithExtraFields):
    source_file: str
    turn_count: typing.Optional[int] = 0
    extra_fields: typing.Optional[dict[str, typing.Any]] = dc.field(
        default_factory=dict
    )

    def as_key(self):
        return (self.source_file,)


@dc.dataclass
class TidyTranscripts(RowWithExtraFields):
    """

    Limitations:

    - Spreadsheet styling is not retained.
    - Word doc styling is not retained.

    """

    transcripts: dict[str, Document]
    split_speaker_on: re.Pattern | str = ":\t"
    spreadsheet_bytes: typing.Optional[bytes] = None

    turns: list[Turn] = dc.field(init=False, default_factory=list)
    segments: list[Segment] = dc.field(init=False, default_factory=list)
    speaker_codes: list[SpeakerCode] = dc.field(init=False, default_factory=list)
    transcript_stats: list[Transcript] = dc.field(init=False, default_factory=list)

    def __post_init__(self):

        if isinstance(self.split_speaker_on, str):
            self.split_speaker_on = re.compile(self.split_speaker_on)

        self.turns = []

        self.speaker_codes = []
        self.segments = []
        self.transcript_stats = []

        speaker_code_counts = Counter()
        segment_counts = Counter()
        transcript_turn_counts = Counter()

        # Extract and populate all the data we need.
        for source_file, doc in self.transcripts.items():
            for turn_details in extract_turns(doc, self.split_speaker_on):

                turn = Turn(source_file, *turn_details)
                self.turns.append(turn)

                # Extract keys for segments and speaker_codes as we go.
                speaker_code_counts[(turn.source_file, turn.speaker_code)] += 1
                segment_counts[(turn.source_file, turn.segment_no)] += 1
                transcript_turn_counts[turn.source_file] += 1

        # Turn them into dataclasses for final processing
        for speaker_code, count in speaker_code_counts.items():
            self.speaker_codes.append(SpeakerCode(*speaker_code, count))

        for segment, count in segment_counts.items():
            self.segments.append(Segment(*segment, count))

        for trans, count in transcript_turn_counts.items():
            self.transcript_stats.append(Transcript(trans, count))

    @classmethod
    def from_filepaths(
        cls, transcript_paths, spreadsheet_path=None, split_speaker_on=":\t"
    ):
        """Load transcripts from files on disk given by paths."""

        transcripts = {}

        for path in transcript_paths:
            with open(path, "rb") as transcript:
                transcripts[path] = Document(transcript)

        spreadsheet_bytes = None

        if spreadsheet_path:
            with open(spreadsheet_path, "rb") as f:
                spreadsheet_bytes = f.read()

        return cls(
            transcripts=transcripts,
            spreadsheet_bytes=spreadsheet_bytes,
            split_speaker_on=split_speaker_on,
        )

    @classmethod
    def from_zip(cls, zip_reader, split_speaker_on=":\t"):
        """Load transcripts from a given zip container."""

        # transcripts = {}

        # for path in transcript_paths:
        #     with open(path, "rb") as transcript:
        #         transcripts[path] = Document(transcript)

        # spreadsheet_bytes = None

        # if spreadsheet_path:
        #     with open(spreadsheet_path, "rb") as f:
        #         spreadsheet_bytes = f.read()

        return cls(
            transcripts=transcripts,
            spreadsheet_bytes=spreadsheet_bytes,
            split_speaker_on=split_speaker_on,
        )

    @classmethod
    def from_ipywidgets(cls, doc_widget, spreadsheet_widget, split_speaker_on=":\t"):
        """Tidy transcripts from the upload widgets"""

        transcripts = {}

        for uploaded in doc_widget.value:
            transcripts[uploaded.name] = Document(uploaded.content)

        spreadsheet_bytes = None

        if spreadsheet_widget.value:
            spreadsheet_bytes = spreadsheet_widget.value[0].content.to_bytes()

        return cls(
            transcripts=transcripts,
            spreadsheet_bytes=spreadsheet_bytes,
            split_speaker_on=split_speaker_on,
        )

    def extract_from_existing_spreadsheet(self):

        if self.spreadsheet_bytes:
            wb = load_workbook(filename=BytesIO(self.spreadsheet_bytes))
        else:
            wb = Workbook()

        speaker_codes = {}
        segments = {}
        transcript_stats = {}

        sheet_map = [
            # This is sheetname, primary key columns, the mapped datatype, and the place
            # to put results. Only extra_fields are preserved: the other column is the
            # turn count which needs to be regenerated.
            (
                "speaker_code",
                ("source_file", "speaker_code"),
                SpeakerCode,
                speaker_codes,
            ),
            ("segment", ("source_file", "segment_no"), Segment, segments),
            ("transcript_file", ("source_file",), Transcript, transcript_stats),
        ]

        for sheetname, key_columns, sheet_type, data_loc in sheet_map:

            if sheetname in wb.sheetnames:

                ws = wb[sheetname]
                rows = ws.iter_rows()

                header_row = next(rows)

                header = [cell.value for cell in header_row]

                if not all(column in header for column in key_columns):
                    raise ValueError(
                        f"{sheetname} needs key columns {key_columns} to be mergable"
                    )

                for row in rows:
                    values = {h: cell.value for h, cell in zip(header, row)}
                    # Always update turn count
                    values.pop("turn_count", None)

                    key = tuple(values[key] for key in key_columns)
                    extra_fields = {
                        key: value
                        for key, value in values.items()
                        if key not in key_columns
                    }

                    if key in data_loc:
                        raise ValueError(
                            f"Repeating key for {sheetname} with value {key}"
                        )

                    data_loc[key] = sheet_type(*key, extra_fields=extra_fields)

        return segments, speaker_codes, transcript_stats

    def as_xlsx(self):
        """
        Return a new xlsx file with all of the data merged together.

        This will create a copy of the original provided spreadsheet, replace the turn
        sheet with the new content, and merge the speaker_code, transcript and segment
        sheets together.

        """

        # Extract existing information in spreadsheet for merging
        (
            existing_segments,
            existing_speaker_codes,
            existing_transcript_stats,
        ) = self.extract_from_existing_spreadsheet()

        if self.spreadsheet_bytes:
            wb = load_workbook(filename=BytesIO(self.spreadsheet_bytes))
        else:
            wb = Workbook()
            wb.remove(wb["Sheet"])

        # Make sure turns sheet exists
        if "turn" in wb.sheetnames:
            wb.remove(wb["turn"])

        turn_sheet = wb.create_sheet("turn")

        header = [f.name for f in dc.fields(Turn)]
        turn_sheet.append(header)

        for turn in self.turns:
            turn_sheet.append(dc.astuple(turn))

        config = [
            ("speaker_code", self.speaker_codes, existing_speaker_codes),
            ("segment", self.segments, existing_segments),
            ("transcript_file", self.transcript_stats, existing_transcript_stats),
        ]
        for sheet_name, rows, extra_data in config:

            if sheet_name in wb.sheetnames:
                wb.remove(wb[sheet_name])

            sheet = wb.create_sheet(sheet_name)

            # Intercept if there are any extra fields...
            # Again, this only works if the extra rows are consistent.
            temp_row = rows[0]
            if extra_data:
                temp_row = next(iter(extra_data.values()))

            sheet.append(temp_row.as_header_row())

            for row in rows:

                key = row.as_key()

                if key in extra_data:
                    row.merge_extras(extra_data[key])
                    del extra_data[key]

                sheet.append(row.as_row())

            # Write any existing rows that might not have been matched to anything.
            for extra_row in extra_data.values():
                sheet.append(row.as_row())

        return wb


if __name__ == "__main__":

    tidied = TidyTranscripts.from_filepaths(
        ["../examples/transcript_format_example.docx"],
        spreadsheet_path="../examples/output.xlsx",
    )

    print(tidied.extract_from_existing_spreadsheet())
    wb = tidied.as_xlsx()
    wb.save("output.xlsx")
